"""Co-Writer backend tests: doc id validation, storage CRUD, history limits."""

from io import BytesIO
from pathlib import Path
import zipfile

from docx import Document as DocxDocument
from fastapi import FastAPI, HTTPException
from fastapi.testclient import TestClient
import pytest

import deeptutor.services.config as _dt_config

_dt_config.load_config_with_main = lambda *_a, **_k: {
    "paths": {},
    "logging": {},
    "system": {"language": "en"},
}

from deeptutor.api.routers import co_writer as co_writer_router
from deeptutor.api.routers.co_writer import _validate_doc_id
from deeptutor.co_writer import edit_agent
from deeptutor.co_writer.docx_converter import (
    DocxConversionError,
    docx_to_markdown,
    markdown_to_docx,
)
from deeptutor.co_writer.storage import CoWriterStorage


class _StubPathService:
    def __init__(self, root: Path):
        self.root = root

    def get_co_writer_dir(self) -> Path:
        return self.root

    def get_co_writer_history_file(self) -> Path:
        return self.root / "history.json"

    def get_co_writer_tool_calls_dir(self) -> Path:
        return self.root / "tool_calls"

    def get_co_writer_docs_dir(self) -> Path:
        return self.root / "documents"

    def get_co_writer_doc_root(self, doc_id: str) -> Path:
        return self.get_co_writer_docs_dir() / f"doc_{doc_id}"

    def get_co_writer_doc_manifest(self, doc_id: str) -> Path:
        return self.get_co_writer_doc_root(doc_id) / "manifest.json"


# ── doc_id validation ────────────────────────────────────────────────────


def test_validate_doc_id_accepts_generated_ids():
    assert _validate_doc_id("a1b2c3d4e5f6") == "a1b2c3d4e5f6"


@pytest.mark.parametrize(
    "bad",
    [
        "../x",
        "a/../../etc",
        "a1b2c3d4e5f6/../../x",
        "doc_1; rm -rf",
        "A1B2C3D4E5F6",
        "",
        "a" * 40,
    ],
)
def test_validate_doc_id_rejects_traversal_and_junk(bad):
    with pytest.raises(HTTPException) as exc:
        _validate_doc_id(bad)
    assert exc.value.status_code == 404


# ── storage CRUD ─────────────────────────────────────────────────────────


def test_storage_crud_roundtrip(tmp_path):
    storage = CoWriterStorage(path_service=_StubPathService(tmp_path))

    doc = storage.create_document(title=None, content="# Hello\nWorld")
    assert doc.title == "Hello"
    assert _validate_doc_id(doc.id) == doc.id

    loaded = storage.load_document(doc.id)
    assert loaded is not None
    assert loaded.content.startswith("# Hello")

    # Explicit titles stick across content updates.
    updated = storage.update_document(doc.id, content="# Renamed\nBody")
    assert updated is not None
    assert updated.title == "Hello"

    assert storage.delete_document(doc.id) is True
    assert storage.load_document(doc.id) is None


def test_storage_untitled_doc_follows_first_heading(tmp_path):
    storage = CoWriterStorage(path_service=_StubPathService(tmp_path))
    doc = storage.create_document(title=None, content="")
    assert doc.title == "Untitled draft"

    updated = storage.update_document(doc.id, content="# Fresh title\nBody")
    assert updated is not None
    assert updated.title == "Fresh title"


def test_storage_list_sorted_by_recency(tmp_path):
    storage = CoWriterStorage(path_service=_StubPathService(tmp_path))
    first = storage.create_document(title="first", content="a")
    second = storage.create_document(title="second", content="b")
    storage.update_document(first.id, content="a updated")

    summaries = storage.list_documents()
    assert [s.id for s in summaries][0] == first.id
    assert {s.id for s in summaries} == {first.id, second.id}


# ── history limits ───────────────────────────────────────────────────────


def test_append_history_caps_entries(tmp_path, monkeypatch):
    stub = _StubPathService(tmp_path)
    monkeypatch.setattr(edit_agent, "get_path_service", lambda: stub)

    overflow = 5
    for i in range(edit_agent._HISTORY_MAX_ENTRIES + overflow):
        edit_agent.append_history({"id": str(i)})

    history = edit_agent.load_history()
    assert len(history) == edit_agent._HISTORY_MAX_ENTRIES
    assert history[-1]["id"] == str(edit_agent._HISTORY_MAX_ENTRIES + overflow - 1)
    assert history[0]["id"] == str(overflow)


def test_append_history_clips_long_texts(tmp_path, monkeypatch):
    stub = _StubPathService(tmp_path)
    monkeypatch.setattr(edit_agent, "get_path_service", lambda: stub)

    long_text = "x" * (edit_agent._HISTORY_TEXT_LIMIT + 500)
    edit_agent.append_history({"id": "clip", "input": {"original_text": long_text}})

    record = edit_agent.load_history()[-1]
    stored = record["input"]["original_text"]
    assert stored.endswith("…[truncated]")
    assert len(stored) < len(long_text)


def test_load_history_survives_corrupt_file(tmp_path, monkeypatch):
    stub = _StubPathService(tmp_path)
    monkeypatch.setattr(edit_agent, "get_path_service", lambda: stub)

    stub.get_co_writer_dir().mkdir(parents=True, exist_ok=True)
    stub.get_co_writer_history_file().write_text("{not json", encoding="utf-8")
    assert edit_agent.load_history() == []


def _docx_bytes_from_document(document: DocxDocument) -> bytes:
    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def _make_docx_with_structure() -> bytes:
    document = DocxDocument()
    document.add_heading("Quarterly Notes", level=1)
    document.add_paragraph("Intro paragraph.")
    document.add_paragraph("First item", style="List Bullet")
    document.add_paragraph("Second item", style="List Bullet")
    document.add_paragraph("Step one", style="List Number")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Revenue"
    table.rows[1].cells[1].text = "1.2M"
    return _docx_bytes_from_document(document)


def test_docx_to_markdown_preserves_structure():
    markdown = docx_to_markdown(_make_docx_with_structure(), "notes.docx")
    assert markdown.startswith("# Quarterly Notes")
    assert "Intro paragraph." in markdown
    assert "- First item" in markdown
    assert "- Second item" in markdown
    assert "1. Step one" in markdown
    assert "| Metric | Value |" in markdown
    assert "| Revenue | 1.2M |" in markdown


def test_markdown_to_docx_roundtrip_reopens():
    markdown = (
        "# Title\n\n"
        "A **bold** and *italic* sentence.\n\n"
        "- Alpha\n\n"
        "- Bravo\n\n"
        "1. First\n\n"
        "| Col | Val |\n"
        "| --- | --- |\n"
        "| A | 1 |\n"
    )
    data = markdown_to_docx(markdown, title="Title")
    restored = DocxDocument(BytesIO(data))
    texts = [p.text for p in restored.paragraphs if p.text.strip()]
    assert "Title" in texts
    assert any("bold" in text and "italic" in text for text in texts)
    assert restored.tables
    assert restored.tables[0].rows[0].cells[0].text == "Col"
    assert restored.tables[0].rows[1].cells[1].text == "1"


def test_docx_to_markdown_rejects_empty_and_corrupt():
    with pytest.raises(DocxConversionError):
        docx_to_markdown(b"", "empty.docx")
    with pytest.raises(DocxConversionError):
        docx_to_markdown(b"not-a-docx", "bad.docx")
    with pytest.raises(DocxConversionError):
        docx_to_markdown(b"\xd0\xcf\x11\xe0" + b"\x00" * 32, "legacy.doc")


def test_docx_to_markdown_rejects_suspicious_zip(monkeypatch):
    monkeypatch.setattr("deeptutor.co_writer.docx_converter._DOCX_MAX_MEMBERS", 2)
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w") as archive:
        archive.writestr("a.xml", "<a/>")
        archive.writestr("b.xml", "<b/>")
        archive.writestr("c.xml", "<c/>")
    with pytest.raises(DocxConversionError, match="too many archive members"):
        docx_to_markdown(buf.getvalue(), "bomb.docx")


def _client(tmp_path, monkeypatch) -> TestClient:
    storage = CoWriterStorage(path_service=_StubPathService(tmp_path))
    monkeypatch.setattr(co_writer_router, "get_co_writer_storage", lambda: storage)
    app = FastAPI()
    app.include_router(co_writer_router.router)
    return TestClient(app), storage


def test_import_docx_creates_document(tmp_path, monkeypatch):
    client, storage = _client(tmp_path, monkeypatch)
    response = client.post(
        "/documents/import/docx",
        files={
            "file": (
                "Quarterly Notes.docx",
                _make_docx_with_structure(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["title"] == "Quarterly Notes"
    assert payload["content"].startswith("# Quarterly Notes")
    loaded = storage.load_document(payload["id"])
    assert loaded is not None
    assert loaded.content.startswith("# Quarterly Notes")


def test_import_docx_rejects_doc_and_unsupported(tmp_path, monkeypatch):
    client, _storage = _client(tmp_path, monkeypatch)
    doc = client.post(
        "/documents/import/docx",
        files={"file": ("legacy.doc", b"OLE", "application/msword")},
    )
    assert doc.status_code == 400
    assert "docx" in doc.json()["detail"].lower()

    txt = client.post(
        "/documents/import/docx",
        files={"file": ("notes.txt", b"hello", "text/plain")},
    )
    assert txt.status_code == 400


def test_export_docx_uses_request_content(tmp_path, monkeypatch):
    client, storage = _client(tmp_path, monkeypatch)
    stored = storage.create_document(title="Stored", content="# Old\nStored body")
    response = client.post(
        "/documents/export/docx",
        json={"title": "Live Title", "content": "# Live\nUnsaved body"},
    )
    assert response.status_code == 200, response.text
    assert (
        response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert 'filename="Live Title.docx"' in response.headers["content-disposition"]
    exported = DocxDocument(BytesIO(response.content))
    texts = [p.text for p in exported.paragraphs if p.text.strip()]
    assert "Live" in texts
    assert "Unsaved body" in texts
    assert "Stored body" not in texts
    assert storage.load_document(stored.id).content.endswith("Stored body")


def test_import_docx_rejects_oversized_upload(tmp_path, monkeypatch):
    monkeypatch.setattr(co_writer_router, "_MAX_DOCX_UPLOAD_BYTES", 16)
    client, _storage = _client(tmp_path, monkeypatch)
    response = client.post(
        "/documents/import/docx",
        files={
            "file": (
                "big.docx",
                b"PK\x03\x04" + b"x" * 32,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 400
    assert "MB limit" in response.json()["detail"]


def _numbered_docx(num_fmt: str) -> bytes:
    """A paragraph carrying a real numPr plus a numbering part in `num_fmt`."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = DocxDocument()
    paragraph = document.add_paragraph("Numbered thing", style="List Paragraph")
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "7")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)

    numbering = document.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), "3")
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl.append(fmt)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "7")
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), "3")
    num.append(ref)
    numbering.append(num)
    return _docx_bytes_from_document(document)


def test_numbered_list_uses_word_numbering_format_not_the_style_name():
    """`List Paragraph` + decimal numbering is an ordered list, not a bullet."""
    assert docx_to_markdown(_numbered_docx("decimal")).startswith("1. ")
    assert docx_to_markdown(_numbered_docx("bullet")).startswith("- ")


def test_hyperlink_text_is_preserved_as_a_markdown_link():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = DocxDocument()
    paragraph = document.add_paragraph("See ")
    r_id = document.part.relate_to(
        "https://example.com",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    node = OxmlElement("w:t")
    node.text = "this link"
    run.append(node)
    link.append(run)
    paragraph._p.append(link)
    paragraph.add_run(" now.")

    markdown = docx_to_markdown(_docx_bytes_from_document(document))
    assert markdown == "See [this link](https://example.com) now."


def test_markdown_link_round_trips_back_into_a_word_hyperlink():
    data = markdown_to_docx("See [this link](https://example.com) now.")
    restored = DocxDocument(BytesIO(data))
    assert restored.paragraphs[0].text == "See this link now."
    assert docx_to_markdown(data) == "See [this link](https://example.com) now."


def test_escaped_literals_survive_a_docx_markdown_docx_round_trip():
    document = DocxDocument()
    document.add_paragraph("cost is 5*3 and a_b_c")
    markdown = docx_to_markdown(_docx_bytes_from_document(document))
    restored = DocxDocument(BytesIO(markdown_to_docx(markdown)))
    assert restored.paragraphs[0].text == "cost is 5*3 and a_b_c"


def test_inline_code_survives_a_round_trip():
    markdown = "Use `foo_bar()` here."
    data = markdown_to_docx(markdown)
    assert DocxDocument(BytesIO(data)).paragraphs[0].text == "Use foo_bar() here."
    assert docx_to_markdown(data) == markdown


def test_table_cells_keep_escaped_pipes_instead_of_splitting():
    markdown = "| a | b |\n| --- | --- |\n| x \\| y | 2 |"
    table = DocxDocument(BytesIO(markdown_to_docx(markdown))).tables[0]
    assert [c.text for c in table.rows[0].cells] == ["a", "b"]
    assert [c.text for c in table.rows[1].cells] == ["x | y", "2"]


def test_headings_round_trip_as_headings():
    data = markdown_to_docx("## Section")
    restored = DocxDocument(BytesIO(data))
    assert restored.paragraphs[0].style.name == "Heading 2"
    assert docx_to_markdown(data) == "## Section"


def test_export_docx_handles_a_non_ascii_title(tmp_path, monkeypatch):
    client, _storage = _client(tmp_path, monkeypatch)
    response = client.post(
        "/documents/export/docx",
        json={"title": "季度报告", "content": "# Hi"},
    )
    assert response.status_code == 200, response.text
    disposition = response.headers["content-disposition"]
    assert "filename*=UTF-8''" in disposition
    assert "%E5%AD%A3" in disposition


def test_importing_the_api_does_not_pull_python_docx_into_memory():
    """The converter must stay import-cheap; python-docx loads on first use."""
    import subprocess
    import sys

    probe = "import sys;import deeptutor.co_writer.docx_converter;print('docx' in sys.modules)"
    result = subprocess.run(
        [sys.executable, "-c", probe],
        cwd=Path(__file__).resolve().parents[2],
        capture_output=True,
        check=True,
        text=True,
    )
    assert result.stdout.strip() == "False"


def test_fenced_code_blocks_round_trip_as_fences_not_inline_code():
    markdown = "```\ndef f():\n    return 1\n```"
    assert docx_to_markdown(markdown_to_docx(markdown)) == markdown


def test_a_lone_monospaced_paragraph_stays_inline_code():
    assert docx_to_markdown(markdown_to_docx("`solo()`")) == "`solo()`"
