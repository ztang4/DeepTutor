"""Markdown <-> DOCX conversion for Co-Writer interchange.

``python-docx`` is imported lazily so importing the API surface does not pull
the OOXML stack into memory; see
``tests/runtime/test_api_import_memory_boundary.py``.
"""

from __future__ import annotations

from collections import deque
from io import BytesIO
import re
from typing import Any, Iterator
import zipfile

_OOXML_MAGIC = b"PK\x03\x04"
_OLE_MAGIC = b"\xd0\xcf\x11\xe0"
_DOCX_MAX_MEMBERS = 4096
_DOCX_MAX_MEMBER_BYTES = 20 * 1024 * 1024
_DOCX_MAX_TOTAL_UNCOMPRESSED = 200 * 1024 * 1024
_DOCX_MAX_COMPRESSION_RATIO = 200.0
_MAX_MARKDOWN_CHARS = 600_000

_HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
_CODE_FONTS = {"consolas", "courier new", "courier", "menlo", "monaco", "cascadia mono"}
_CODE_FONT = "Consolas"
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

_HEADING_STYLES = {
    "title": 1,
    "heading 1": 1,
    "heading 2": 2,
    "heading 3": 3,
    "heading 4": 4,
    "heading 5": 5,
    "heading 6": 6,
    "subtitle": 2,
}
_QUOTE_STYLES = {"quote", "intense quote", "blockquote"}

# Word numbering formats that should render as an ordered list.
_ORDERED_NUM_FMTS = {
    "decimal",
    "decimalzero",
    "lowerletter",
    "upperletter",
    "lowerroman",
    "upperroman",
    "ordinal",
    "ordinaltext",
    "cardinaltext",
    "decimalenclosedcircle",
    "decimalenclosedparen",
}

_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_MD_UL_RE = re.compile(r"^[-*+]\s+(.*)$")
_MD_OL_RE = re.compile(r"^\d+[.)]\s+(.*)$")
_MD_QUOTE_RE = re.compile(r"^>\s?(.*)$")
_MD_HR_RE = re.compile(r"^(-{3,}|\*{3,}|_{3,})$")
_MD_TABLE_SEP_RE = re.compile(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)*\|?\s*$")
_MD_FENCE_RE = re.compile(r"^```")
_INLINE_RE = re.compile(
    r"\[(?P<link_text>[^\]\n]*)\]\((?P<link_href>[^)\s]*)\)"
    r"|`(?P<code>[^`\n]+)`"
    r"|\*\*\*(?P<bi>[^\n]+?)\*\*\*"
    r"|\*\*(?P<b1>[^\n]+?)\*\*"
    r"|__(?P<b2>[^\n]+?)__"
    r"|~~(?P<strike>[^\n]+?)~~"
    r"|\*(?P<i1>[^\n]+?)\*"
    r"|_(?P<i2>[^\n]+?)_"
)
# Characters we escape when emitting markdown, and honour when parsing it.
_ESCAPABLE = set("\\`*_[]~|<>#+-.!()")
_ESCAPE_RE = re.compile(r"([\\`*_\[\]~|])")
# A "|" not preceded by an odd number of backslashes.
_UNESCAPED_PIPE_RE = re.compile(r"(?<!\\)(?:\\\\)*\|")
_MASK = "\x00"


class DocxConversionError(Exception):
    pass


def _docx_module() -> Any:
    """Import ``docx`` on demand, translating a missing dependency."""
    try:
        import docx
        import docx.oxml
        import docx.oxml.ns
        import docx.table
        import docx.text.paragraph
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise DocxConversionError(
            "Word support is unavailable because python-docx is not installed."
        ) from exc
    return docx


# --------------------------------------------------------------------------
# DOCX -> Markdown
# --------------------------------------------------------------------------


def docx_to_markdown(data: bytes, filename: str = "document.docx") -> str:
    if not data:
        raise DocxConversionError(f"{filename} is empty.")
    if data.startswith(_OLE_MAGIC):
        raise DocxConversionError(
            "Legacy .doc files are not supported yet. "
            "Please save the document as .docx and try again."
        )
    if not data.startswith(_OOXML_MAGIC):
        raise DocxConversionError(
            f"{filename} is not a valid Word document. "
            "Please open it in Word and save it as a new .docx file."
        )
    _validate_docx_archive(data, filename)
    try:
        docx = _docx_module()
        document = docx.Document(BytesIO(data))
        numbering = _numbering_formats(document)
        blocks = [_block_to_markdown(block, numbering) for block in _iter_blocks(document)]
        markdown = _join_blocks(blocks)
    except DocxConversionError:
        raise
    except Exception as exc:
        markdown = _plain_ooxml_fallback(data, filename, exc)
    cleaned = markdown.strip()
    if not cleaned:
        raise DocxConversionError(f"{filename} has no extractable text.")
    if len(cleaned) > _MAX_MARKDOWN_CHARS:
        raise DocxConversionError(
            f"{filename} exceeds the Co-Writer character limit "
            f"({_MAX_MARKDOWN_CHARS:,} characters)."
        )
    return cleaned


def _validate_docx_archive(data: bytes, filename: str) -> None:
    try:
        archive = zipfile.ZipFile(BytesIO(data))
    except zipfile.BadZipFile as exc:
        raise DocxConversionError(
            f"{filename} is not a valid Word document. "
            "Please open it in Word and save it as a new .docx file."
        ) from exc
    with archive:
        members = [info for info in archive.infolist() if not info.is_dir()]
        if len(members) > _DOCX_MAX_MEMBERS:
            raise DocxConversionError(f"{filename} has too many archive members ({len(members)}).")
        total = 0
        for info in members:
            if info.file_size > _DOCX_MAX_MEMBER_BYTES:
                raise DocxConversionError(
                    f"{filename}: archive member {info.filename} is too large."
                )
            total += info.file_size
            if total > _DOCX_MAX_TOTAL_UNCOMPRESSED:
                raise DocxConversionError(f"{filename}: uncompressed contents are too large.")
            if (
                info.compress_size
                and info.file_size / info.compress_size > _DOCX_MAX_COMPRESSION_RATIO
            ):
                raise DocxConversionError(
                    f"{filename}: archive member {info.filename} "
                    "has a suspicious compression ratio."
                )


def _plain_ooxml_fallback(data: bytes, filename: str, cause: Exception) -> str:
    from deeptutor.utils.document_extractor import extract_text_from_bytes

    try:
        return extract_text_from_bytes(filename, data, max_bytes=None, max_chars=None)
    except Exception as exc:
        raise DocxConversionError(
            f"{filename} could not be read. Please open it in Word and save it as a new .docx file."
        ) from (exc or cause)


def _numbering_formats(document: Any) -> dict[tuple[str, str], str]:
    """Map ``(numId, ilvl)`` to the Word numbering format, when available."""
    docx = _docx_module()
    qn = docx.oxml.ns.qn
    try:
        numbering = document.part.numbering_part.element
    except Exception:
        return {}
    try:
        abstract: dict[str, dict[str, str]] = {}
        for node in numbering.findall(qn("w:abstractNum")):
            abstract_id = node.get(qn("w:abstractNumId"))
            if abstract_id is None:
                continue
            levels: dict[str, str] = {}
            for lvl in node.findall(qn("w:lvl")):
                ilvl = lvl.get(qn("w:ilvl")) or "0"
                fmt_node = lvl.find(qn("w:numFmt"))
                if fmt_node is not None:
                    levels[ilvl] = (fmt_node.get(qn("w:val")) or "").lower()
            abstract[abstract_id] = levels
        formats: dict[tuple[str, str], str] = {}
        for node in numbering.findall(qn("w:num")):
            num_id = node.get(qn("w:numId"))
            ref = node.find(qn("w:abstractNumId"))
            if num_id is None or ref is None:
                continue
            for ilvl, fmt in abstract.get(ref.get(qn("w:val")) or "", {}).items():
                formats[(num_id, ilvl)] = fmt
        return formats
    except Exception:
        return {}


def _iter_blocks(document: Any) -> Iterator[Any]:
    docx = _docx_module()
    qn = docx.oxml.ns.qn
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield docx.text.paragraph.Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield docx.table.Table(child, document)


def _block_to_markdown(block: Any, numbering: dict[tuple[str, str], str]) -> str:
    if type(block).__name__ == "Table":
        return _table_to_markdown(block)
    return _paragraph_to_markdown(block, numbering)


def _paragraph_to_markdown(paragraph: Any, numbering: dict[tuple[str, str], str]) -> str:
    if not (paragraph.text or "").strip():
        return ""
    style_name = _style_name(paragraph)
    lower = style_name.lower()
    body = _inline_content_to_markdown(paragraph) or _escape_md(paragraph.text.strip())
    heading_level = _HEADING_STYLES.get(lower)
    if heading_level:
        return f"{'#' * heading_level} {body}"
    kind, level = _list_kind(paragraph, lower, numbering)
    indent = "  " * level
    if kind == "ol":
        return f"{indent}1. {body}"
    if kind == "ul":
        return f"{indent}- {body}"
    if lower in _QUOTE_STYLES:
        return f"> {body}"
    return body


def _style_name(paragraph: Any) -> str:
    try:
        if paragraph.style is not None and paragraph.style.name:
            return str(paragraph.style.name)
    except Exception:
        pass
    return ""


def _list_kind(
    paragraph: Any, lower_style: str, numbering: dict[tuple[str, str], str]
) -> tuple[str | None, int]:
    """Return ``(kind, indent_level)`` where kind is ``ul``, ``ol`` or ``None``."""
    if "heading" in lower_style or lower_style in {"title", "subtitle"}:
        return None, 0
    num_id, ilvl = _numbering_ref(paragraph)
    level = int(ilvl) if ilvl and ilvl.isdigit() else 0
    level = min(level, 5)
    if num_id is not None:
        fmt = numbering.get((num_id, ilvl or "0"), "")
        if fmt == "bullet":
            return "ul", level
        if fmt in _ORDERED_NUM_FMTS:
            return "ol", level
    # Fall back to the style name when numbering.xml is absent or unreadable.
    if "number" in lower_style or "enum" in lower_style:
        return "ol", level
    if "bullet" in lower_style:
        return "ul", level
    if num_id is not None:
        return "ul", level
    return None, 0


def _numbering_ref(paragraph: Any) -> tuple[str | None, str | None]:
    docx = _docx_module()
    qn = docx.oxml.ns.qn
    try:
        p_pr = paragraph._p.pPr
        if p_pr is None or p_pr.numPr is None:
            return None, None
        num_pr = p_pr.numPr
        num_id_node = num_pr.find(qn("w:numId"))
        ilvl_node = num_pr.find(qn("w:ilvl"))
        num_id = num_id_node.get(qn("w:val")) if num_id_node is not None else None
        ilvl = ilvl_node.get(qn("w:val")) if ilvl_node is not None else "0"
        return num_id, ilvl
    except Exception:
        return None, None


def _inline_content_to_markdown(paragraph: Any) -> str:
    """Render runs *and* hyperlinks; ``paragraph.runs`` omits linked runs."""
    parts: list[str] = []
    try:
        items = list(paragraph.iter_inner_content())
    except Exception:
        items = list(paragraph.runs)
    for item in items:
        if type(item).__name__ == "Hyperlink":
            text = _runs_to_markdown(getattr(item, "runs", [])) or _escape_md(item.text or "")
            if not text:
                continue
            address = (getattr(item, "address", "") or "").strip()
            parts.append(f"[{text}]({address})" if address else text)
        else:
            parts.append(_run_to_markdown(item))
    return "".join(parts)


def _runs_to_markdown(runs: Any) -> str:
    return "".join(_run_to_markdown(run) for run in runs)


def _run_to_markdown(run: Any) -> str:
    text = run.text or ""
    if not text:
        return ""
    try:
        font_name = (run.font.name or "").lower()
    except Exception:
        font_name = ""
    if font_name in _CODE_FONTS:
        return f"`{text}`" if "`" not in text else _escape_md(text)
    out = _escape_md(text)
    try:
        if run.font.strike:
            out = f"~~{out}~~"
    except Exception:
        pass
    if run.bold and run.italic:
        return f"***{out}***"
    if run.bold:
        return f"**{out}**"
    if run.italic:
        return f"*{out}*"
    return out


def _table_to_markdown(table: Any) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        try:
            cells = [_cell_text(cell) for cell in row.cells]
        except Exception:
            continue
        if any(cell.strip() for cell in cells):
            rows.append(cells)
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _cell_text(cell: Any) -> str:
    text = " ".join(part.strip() for part in (cell.text or "").splitlines() if part.strip())
    return text.replace("\\", "\\\\").replace("|", "\\|")


def _join_blocks(blocks: list[str]) -> str:
    return "\n\n".join(_group_code_blocks(blocks))


def _group_code_blocks(blocks: list[str]) -> list[str]:
    """Re-fence runs of monospaced paragraphs that were a code block.

    A single monospaced paragraph stays inline code; two or more in a row are
    far more likely to be a fenced block that ``markdown_to_docx`` flattened.
    """
    out: list[str] = []
    run: list[str] = []

    def flush() -> None:
        if not run:
            return
        if len(run) == 1:
            out.append(run[0])
        else:
            out.append("```\n" + "\n".join(line[1:-1] for line in run) + "\n```")
        run.clear()

    for block in blocks:
        if not block.strip():
            continue
        if _is_code_line(block):
            run.append(block)
            continue
        flush()
        out.append(block)
    flush()
    return out


def _is_code_line(block: str) -> bool:
    return (
        len(block) >= 2
        and block.startswith("`")
        and block.endswith("`")
        and "`" not in block[1:-1]
        and "\n" not in block
    )


def _escape_md(text: str) -> str:
    return _ESCAPE_RE.sub(r"\\\1", text)


# --------------------------------------------------------------------------
# Markdown -> DOCX
# --------------------------------------------------------------------------


def markdown_to_docx(content: str, title: str = "") -> bytes:
    del title  # any heading already lives in the markdown body
    docx = _docx_module()
    document = docx.Document()
    lines = (content or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if _MD_FENCE_RE.match(stripped):
            index = _append_code_block(document, lines, index)
            continue
        if (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and _MD_TABLE_SEP_RE.match(lines[index + 1].strip())
        ):
            index = _append_markdown_table(document, lines, index)
            continue
        heading = _MD_HEADING_RE.match(stripped)
        if heading:
            level = min(len(heading.group(1)), 6)
            text = heading.group(2).strip()
            if text:
                _append_styled_paragraph(document, text, f"Heading {level}", heading_level=level)
            index += 1
            continue
        if _MD_HR_RE.match(stripped):
            index += 1
            continue
        quote = _MD_QUOTE_RE.match(stripped)
        if quote:
            _append_styled_paragraph(document, quote.group(1), "Quote")
            index += 1
            continue
        indent = _leading_indent(line)
        unordered = _MD_UL_RE.match(stripped)
        if unordered:
            _append_styled_paragraph(
                document, unordered.group(1), _list_style("List Bullet", indent)
            )
            index += 1
            continue
        ordered = _MD_OL_RE.match(stripped)
        if ordered:
            _append_styled_paragraph(document, ordered.group(1), _list_style("List Number", indent))
            index += 1
            continue
        paragraph = document.add_paragraph()
        _add_inline_runs(paragraph, stripped)
        index += 1
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _leading_indent(line: str) -> int:
    expanded = line.replace("\t", "    ")
    return min((len(expanded) - len(expanded.lstrip(" "))) // 2, 2)


def _list_style(base: str, indent: int) -> str:
    return base if indent <= 0 else f"{base} {indent + 1}"


def _append_styled_paragraph(
    document: Any, text: str, style: str, *, heading_level: int | None = None
) -> None:
    try:
        paragraph = document.add_paragraph(style=style)
    except (KeyError, ValueError):
        if heading_level is not None:
            paragraph = document.add_heading("", level=heading_level)
        else:
            paragraph = document.add_paragraph()
    _add_inline_runs(paragraph, text)


def _append_code_block(document: Any, lines: list[str], index: int) -> int:
    index += 1
    collected: list[str] = []
    while index < len(lines):
        if _MD_FENCE_RE.match(lines[index].strip()):
            index += 1
            break
        collected.append(lines[index])
        index += 1
    for line in collected or [""]:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = _CODE_FONT
    return index


def _append_markdown_table(document: Any, lines: list[str], index: int) -> int:
    rows: list[list[str]] = [_split_table_row(lines[index])]
    index += 2
    while index < len(lines) and lines[index].strip().startswith("|"):
        rows.append(_split_table_row(lines[index]))
        index += 1
    cols = max((len(row) for row in rows), default=0)
    if cols == 0:
        return index
    table = document.add_table(rows=len(rows), cols=cols)
    try:
        table.style = "Table Grid"
    except (KeyError, ValueError):
        pass
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            table.rows[r_idx].cells[c_idx].text = row[c_idx] if c_idx < len(row) else ""
    return index


def _split_table_row(line: str) -> list[str]:
    r"""Split on unescaped ``|`` so that ``\|`` stays inside a cell."""
    stripped = line.strip()
    cells: list[str] = []
    start = 0
    for match in _UNESCAPED_PIPE_RE.finditer(stripped):
        pos = match.end() - 1
        cells.append(stripped[start:pos])
        start = pos + 1
    cells.append(stripped[start:])
    if cells and not cells[0].strip():
        cells.pop(0)
    if cells and not cells[-1].strip():
        cells.pop()
    return [_unescape_md(cell.strip()) for cell in cells]


def _unescape_md(text: str) -> str:
    out: list[str] = []
    i = 0
    while i < len(text):
        if text[i] == "\\" and i + 1 < len(text) and text[i + 1] in _ESCAPABLE:
            out.append(text[i + 1])
            i += 2
        else:
            out.append(text[i])
            i += 1
    return "".join(out)


def _mask_escapes(text: str) -> tuple[str, deque[str]]:
    r"""Replace ``\X`` escapes with a sentinel so markers cannot match them."""
    out: list[str] = []
    saved: deque[str] = deque()
    i = 0
    while i < len(text):
        if text[i] == "\\" and i + 1 < len(text) and text[i + 1] in _ESCAPABLE:
            saved.append(text[i + 1])
            out.append(_MASK)
            i += 2
        else:
            out.append(text[i])
            i += 1
    return "".join(out), saved


def _unmask(fragment: str, saved: deque[str]) -> str:
    if _MASK not in fragment:
        return fragment
    return "".join(saved.popleft() if ch == _MASK else ch for ch in fragment)


def _add_inline_runs(paragraph: Any, text: str) -> None:
    masked, saved = _mask_escapes(text)
    cursor = 0
    for match in _INLINE_RE.finditer(masked):
        if match.start() > cursor:
            paragraph.add_run(_unmask(masked[cursor : match.start()], saved))
        groups = match.groupdict()
        if groups["link_text"] is not None:
            _add_hyperlink(
                paragraph,
                _unmask(groups["link_text"], saved),
                _unmask(groups["link_href"] or "", saved),
            )
        elif groups["code"] is not None:
            run = paragraph.add_run(_unmask(groups["code"], saved))
            run.font.name = _CODE_FONT
        elif groups["bi"] is not None:
            run = paragraph.add_run(_unmask(groups["bi"], saved))
            run.bold = True
            run.italic = True
        elif groups["b1"] is not None or groups["b2"] is not None:
            run = paragraph.add_run(_unmask(groups["b1"] or groups["b2"] or "", saved))
            run.bold = True
        elif groups["strike"] is not None:
            run = paragraph.add_run(_unmask(groups["strike"], saved))
            run.font.strike = True
        else:
            run = paragraph.add_run(_unmask(groups["i1"] or groups["i2"] or "", saved))
            run.italic = True
        cursor = match.end()
    if cursor < len(masked):
        paragraph.add_run(_unmask(masked[cursor:], saved))


def _add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    if not text:
        return
    if not url:
        paragraph.add_run(text)
        return
    docx = _docx_module()
    qn = docx.oxml.ns.qn
    try:
        r_id = paragraph.part.relate_to(url, _HYPERLINK_REL, is_external=True)
        link = docx.oxml.OxmlElement("w:hyperlink")
        link.set(qn("r:id"), r_id)
        run = docx.oxml.OxmlElement("w:r")
        r_pr = docx.oxml.OxmlElement("w:rPr")
        style = docx.oxml.OxmlElement("w:rStyle")
        style.set(qn("w:val"), "Hyperlink")
        r_pr.append(style)
        run.append(r_pr)
        node = docx.oxml.OxmlElement("w:t")
        node.text = text
        node.set(_XML_SPACE, "preserve")
        run.append(node)
        link.append(run)
        paragraph._p.append(link)
    except Exception:
        paragraph.add_run(text)


__all__ = [
    "DocxConversionError",
    "docx_to_markdown",
    "markdown_to_docx",
]
